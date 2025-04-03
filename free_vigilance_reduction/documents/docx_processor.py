import docx
from .base import Document

class DocProcessor(Document):
    def get_text(self):
        doc = docx.Document(self.file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    def create_redacted_copy(self, reduced_text):
        doc = docx.Document()
        for paragraph in reduced_text.split('\n'):
            doc.add_paragraph(paragraph)
        output_path = self.file_path.replace('.docx', '_redacted.docx')
        doc.save(output_path)
        return output_path